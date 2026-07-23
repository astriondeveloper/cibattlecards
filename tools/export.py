"""Exports: PDF via headless Chromium, DOCX via python-docx."""
from __future__ import annotations

import os
import re
import shutil
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import build as build_mod
from . import content as c

EXPORT_DIR = c.ROOT / "exports"
CHROME_CANDIDATES = ["chromium", "chromium-browser", "google-chrome", "chrome"]
FALLBACK_CHROME = "/opt/pw-browsers/chromium"


def find_chromium() -> str:
    if os.environ.get("CHROME_BIN"):
        return os.environ["CHROME_BIN"]
    for name in CHROME_CANDIDATES:
        found = shutil.which(name)
        if found:
            return found
    if os.path.exists(FALLBACK_CHROME):
        return FALLBACK_CHROME
    raise SystemExit("No Chromium/Chrome found for PDF export. Install one or "
                     "set CHROME_BIN to its path.")


def _get_card(slug: str):
    for card in c.load_cards():
        if card.slug == slug:
            return card
    raise SystemExit(f"no competitor card with slug '{slug}'")


def export_pdf(slug=None, share_safe=False, everything=False) -> None:
    build_mod.build(quiet=True)
    config = c.load_config()
    docs = config.output_dir
    chrome = find_chromium()

    targets = []
    slugs = ([card.slug for card in c.load_cards()] if everything
             else [_get_card(slug).slug])
    for s in slugs:
        if share_safe:
            targets.append((f"{s}-share-safe.pdf", docs / "share-safe" / f"{s}.html"))
        else:
            targets.append((f"{s}.pdf", docs / "cards" / f"{s}.html"))

    out_dir = EXPORT_DIR / "pdf"
    out_dir.mkdir(parents=True, exist_ok=True)
    for name, page in targets:
        out = out_dir / name
        cmd = [chrome, "--headless=new", "--disable-gpu", "--no-sandbox",
               "--no-pdf-header-footer", f"--print-to-pdf={out}",
               page.resolve().as_uri()]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:  # older Chromium: retry legacy headless
            cmd[1] = "--headless"
            result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise SystemExit(f"PDF export failed for {page.name}:\n"
                             f"{result.stderr.strip()[-800:]}")
        print(f"wrote {out.relative_to(c.ROOT)}")


def _marked_doc(config, title: str, share_safe: bool = False) -> Document:
    doc = Document()
    banner = config.settings["share_safe_banner" if share_safe else "banner"]
    section = doc.sections[0]
    for part in (section.header, section.footer):
        para = part.paragraphs[0]
        para.text = banner
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].bold = True
    created_by = config.settings.get("created_by")
    if created_by:
        credit = section.footer.add_paragraph(f"Created by {created_by}")
        credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(title, level=0)
    return doc


def _add_markdown(doc: Document, text: str) -> None:
    cleaned = c.COMMENT_RE.sub("", text or "").strip()
    for raw in cleaned.splitlines():
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", raw.rstrip())
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.lstrip().startswith("- "):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def export_docx(slug: str, share_safe: bool = False) -> None:
    config = c.load_config()
    card = _get_card(slug)
    records = c.load_winloss()

    doc = _marked_doc(config, card.name, share_safe)
    meta = card.meta
    doc.add_paragraph(
        f"Tier {card.tier} · {meta.get('relationship')} · "
        f"markets: {', '.join(meta.get('markets') or []) or '—'} · "
        f"agencies: {', '.join(meta.get('agencies') or []) or '—'} · "
        f"last reviewed: {meta.get('last_reviewed')}")
    if meta.get("example"):
        doc.add_paragraph("EXAMPLE — fictional demonstration content.")

    omit = set(config.settings.get("share_safe_omit") or []) if share_safe else set()
    for name in c.CARD_SECTIONS:
        if name not in card.sections or name in omit:
            continue
        body = c.visible_text(card.sections[name])
        if name == "Win/Loss vs Us" and not share_safe:
            doc.add_heading(name, level=1)
            tally = c.rollup(records, card.slug)["tally"]
            doc.add_paragraph(
                f"Head-to-head: {tally['pursuits']} pursuits — we won "
                f"{tally['we_won']}, we lost {tally['we_lost']}"
                + (f", {tally['observed']} observed" if tally["observed"] else "")
                + ".")
            if body:
                _add_markdown(doc, body)
            continue
        if not body:
            continue
        doc.add_heading(name, level=1)
        _add_markdown(doc, body)

    out_dir = EXPORT_DIR / "docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{slug}{'-share-safe' if share_safe else ''}.docx"
    doc.save(out)
    print(f"wrote {out.relative_to(c.ROOT)}")


def export_pack(slug: str) -> None:
    """Ghosting & discriminator pack for proposal writers. Internal only."""
    config = c.load_config()
    card = _get_card(slug)
    doc = _marked_doc(config, f"{card.name} — Ghosting & Discriminator Pack")
    doc.add_paragraph("Internal proposal-team material. Ghost language is "
                      "proposal-safe (never names the competitor) but this "
                      "pack itself is not shareable outside the team.")
    if card.meta.get("example"):
        doc.add_paragraph("EXAMPLE — fictional demonstration content.")

    for name in ("Their Discriminators", "Our Counter-Positioning",
                 "Objection Handling"):
        body = c.visible_text(card.sections.get(name, ""))
        if body:
            doc.add_heading(name, level=1)
            _add_markdown(doc, body)

    entries = c.parse_ghost_entries(card.sections.get("Ghosting Library", ""))
    if entries:
        doc.add_heading("Ghosting Library", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, head in enumerate(("Theme", "Ghost language", "Targets", "Proof")):
            table.rows[0].cells[i].text = head
        for entry in entries:
            row = table.add_row().cells
            row[0].text = f"{entry['id']} — {entry['title']}"
            row[1].text = entry["ghost"]
            row[2].text = entry["targets"]
            row[3].text = entry["proof"]

    body = c.visible_text(card.sections.get("Intel & Proof Points", ""))
    if body:
        doc.add_heading("Intel & Proof Points", level=1)
        _add_markdown(doc, body)

    out_dir = EXPORT_DIR / "docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{slug}-pack.docx"
    doc.save(out)
    print(f"wrote {out.relative_to(c.ROOT)}")
